from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


def remove_body_content(doc: Document) -> None:
    body = doc._element.body
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)


def set_cell_margins(cell, top=100, start=110, bottom=100, end=110):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_table_borders(table, color="A9BDD2", size="4"):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), size)
        tag.set(qn("w:color"), color)


def set_table_geometry(table, widths):
    total = sum(widths)
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(total))
    tbl_w.set(qn("w:type"), "dxa")
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        cant_split = OxmlElement("w:cantSplit")
        row._tr.get_or_add_trPr().append(cant_split)
        for cell, width in zip(row.cells, widths):
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def style_table(table, widths):
    set_table_geometry(table, widths)
    set_table_borders(table)
    header = table.rows[0]
    repeat = OxmlElement("w:tblHeader")
    header._tr.get_or_add_trPr().append(repeat)
    for cell in header.cells:
        set_shading(cell, "D9E5F2")
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(31, 77, 120)
    for row in table.rows[1:]:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(0)


def fill_table(doc, rows, widths):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    for row_obj, values in zip(table.rows, rows):
        for cell, value in zip(row_obj.cells, values):
            cell.text = value
    style_table(table, widths)
    return table


def add_title(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run("华中杯 A题｜三人团队启动分工")
    r.bold = True
    r.font.size = Pt(24)
    r.font.color.rgb = RGBColor(22, 58, 95)
    sub = doc.add_paragraph()
    sub.paragraph_format.space_after = Pt(10)
    r = sub.add_run("3天时间 · 只保留最开始怎么着手")
    r.italic = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(107, 114, 128)


def add_callout(doc, label: str, text: str, fill="FFF7E6", accent="D38A19"):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_shading(cell, fill)
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:color"), accent)
    borders.append(left)
    tc_pr.append(borders)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(label + "  ")
    run.bold = True
    run.font.color.rgb = RGBColor(166, 99, 0)
    p.add_run(text)
    return table


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def build(src: Path, dst: Path) -> None:
    doc = Document(src)
    remove_body_content(doc)
    section = doc.sections[0]
    section.different_first_page_header_footer = True

    add_title(doc)
    add_callout(
        doc,
        "共同目标",
        "前6小时先把路线评估器做对：给定车辆和路径，能正确输出到达时间、载荷、能耗、碳排放与五项成本。",
    )

    doc.add_paragraph("1. 三人怎么分工", style="Heading 1")
    fill_table(
        doc,
        [
            ["角色", "主要负责", "最先交付"],
            ["1号：优化实现", "定义解结构和评估器接口；完成初始解、局部搜索与后续求解程序。", "可运行代码骨架与人工样例接口"],
            ["2号：模型验证", "维护模型口径；整理目标与约束；手算路线并编写评估器测试。", "model_spec与1/2客户手算用例"],
            ["3号：数据与整合", "读取附件、审计缺失与异常、聚合并拆分需求；维护实验记录和论文骨架。", "唯一processed数据与data_audit"],
        ],
        [1900, 4100, 3360],
    )

    doc.add_paragraph("1.1 配合规则", style="Heading 2")
    add_bullets(
        doc,
        [
            "3号只生成一套清洗数据；1号不得自行再聚合，2号负责审核处理口径。",
            "2号规定评估器怎么算，1号负责把它实现得稳定、快速；手算与程序必须逐项一致。",
            "所有路线、图表和问题一/二对比统一调用同一个 evaluator.py。",
            "论文由3号整合，1号提供算法说明，2号提供模型与假设，三人共同核对最终数字。",
        ],
    )
    add_callout(doc, "不要重复劳动", "原始Excel只读；成本只算一套；先做对评估器，再写ALNS。", fill="FDEEEE", accent="B91C1C")

    doc.add_page_break()

    doc.add_paragraph("2. 第一天怎么开始", style="Heading 1")
    fill_table(
        doc,
        [
            ["时间", "负责人", "要做什么", "交付"],
            ["0-0.5小时", "全员", "读题、填角色、统一单位和项目目录。", "角色与文件结构"],
            ["0.5-3小时", "三人并行", "1号搭接口；2号写口径和手算；3号完成数据审计与清洗。", "代码骨架+口径+processed数据"],
            ["3-6小时", "1号+2号", "实现并验收 evaluator.py；3号建立结果格式和论文骨架。", "手算与程序一致"],
            ["6-10小时", "全员", "构造问题一首个可行方案，检查漏单、重复、容量和时间递推。", "第一条可行调度方案"],
        ],
        [1450, 1650, 4200, 2060],
    )

    doc.add_paragraph("2.1 开始前必须定下的5件事", style="Heading 2")
    add_bullets(
        doc,
        [
            "缺失重量和体积如何补齐，并保留什么敏感性方案。",
            "绿色区按坐标半径计算，如何处理与题面客户数量不一致。",
            "大客户与超大订单如何拆分，每次到访是否计20分钟服务。",
            "17:00以后车速如何处理，跨交通时段如何分段计算。",
            "问题三默认从问题二合规方案继续动态重调度。",
        ],
    )

    doc.add_paragraph("3. 三天大致节奏", style="Heading 1")
    fill_table(
        doc,
        [
            ["日期", "核心任务", "最低验收"],
            ["第1天", "统一数据口径，完成评估器和问题一首个可行解。", "评估器手算一致、方案可行"],
            ["第2天", "优化问题一；加入绿色区限行并完成问题二对比。", "Q1/Q2路线与成本表"],
            ["第3天", "模拟2类动态事件；整理三问结果并完成论文。", "三问结果齐全、数字一致"],
        ],
        [1450, 4300, 3610],
    )

    dst.parent.mkdir(parents=True, exist_ok=True)
    doc.save(dst)
    print(f"Saved: {dst}")


if __name__ == "__main__":
    if len(sys.argv) != 3:
        raise SystemExit("Usage: build_two_page_team_brief.py INPUT.docx OUTPUT.docx")
    build(Path(sys.argv[1]), Path(sys.argv[2]))
