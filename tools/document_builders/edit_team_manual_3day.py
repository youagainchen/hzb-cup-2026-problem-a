from __future__ import annotations

import sys
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH


def all_paragraphs(doc):
    for paragraph in doc.paragraphs:
        yield paragraph
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph
    for section in doc.sections:
        for part in (section.header, section.footer):
            for paragraph in part.paragraphs:
                yield paragraph


def replace_span(paragraph, old: str, new: str) -> bool:
    runs = paragraph.runs
    full = "".join(run.text for run in runs)
    start = full.find(old)
    if start < 0:
        return False
    end = start + len(old)
    positions = []
    cursor = 0
    for idx, run in enumerate(runs):
        positions.append((idx, cursor, cursor + len(run.text)))
        cursor += len(run.text)
    start_info = next(item for item in positions if item[1] <= start < item[2])
    end_info = next(item for item in positions if item[1] < end <= item[2])
    si, ss, _ = start_info
    ei, es, _ = end_info
    prefix = runs[si].text[: start - ss]
    suffix = runs[ei].text[end - es :]
    runs[si].text = prefix + new + suffix
    for idx in range(si + 1, ei + 1):
        runs[idx].text = ""
    return True


def replace_everywhere(doc, old: str, new: str) -> int:
    count = 0
    for paragraph in all_paragraphs(doc):
        if replace_span(paragraph, old, new):
            count += 1
    return count


def set_paragraph(paragraph, text: str) -> None:
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def replace_exact_paragraph(doc, old: str, new: str) -> None:
    matches = [p for p in all_paragraphs(doc) if p.text.strip() == old]
    if not matches:
        raise ValueError(f"Paragraph not found: {old}")
    for paragraph in matches:
        set_paragraph(paragraph, new)


def table_key(table):
    return tuple(cell.text.strip() for cell in table.rows[0].cells)


def find_table(doc, headers):
    headers = tuple(headers)
    for table in doc.tables:
        if table_key(table) == headers:
            return table
    raise ValueError(f"Table not found: {headers}")


def resize_table(table, target_rows: int) -> None:
    while len(table.rows) > target_rows:
        table._tbl.remove(table.rows[-1]._tr)
    while len(table.rows) < target_rows:
        table._tbl.append(deepcopy(table.rows[-1]._tr))


def write_table(table, rows) -> None:
    resize_table(table, len(rows))
    for row_obj, values in zip(table.rows, rows):
        if len(row_obj.cells) != len(values):
            raise ValueError("Column count mismatch")
        for cell, value in zip(row_obj.cells, values):
            cell.text = str(value)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.keep_together = True
                if len(value) <= 14:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER


def insert_blocks_after(doc, anchor_element, blocks) -> None:
    current = anchor_element
    for text, style in blocks:
        paragraph = doc.add_paragraph(style=style)
        paragraph.add_run(text)
        current.addnext(paragraph._p)
        current = paragraph._p


def main(src: Path, dst: Path) -> None:
    doc = Document(src)

    # Cover and global scale changes.
    replacements = {
        "从第一个文件夹到三问完整结果的可执行起步方案": "从数据清洗到三问完整结果的三天冲刺执行方案",
        "三个人第一天只追求一个共同里程碑": "三个人前6小时只追求一个共同里程碑",
        "第一周个人交付": "三天个人交付",
        "完整实例至少运行10个随机种子": "完整实例至少运行5个随机种子；时间允许再扩展至10个",
        "最终方案进行500-1000次随机速度蒙特卡洛检验": "最终候选方案进行200-500次随机速度蒙特卡洛检验；时间允许扩展至1000次",
        "用Gurobi求10、20、30客户小规模实例，计算ALNS相对误差": "用Gurobi求8、12客户精确解；15客户设置限时并报告最优间隙",
        "最终候选方案进行500-1000次正态速度抽样": "最终候选方案进行200-500次正态速度抽样；时间允许扩展至1000次",
        "至少10个随机种子，报告最好值、均值、标准差和运行时间": "至少5个随机种子，报告最好值、均值、标准差和运行时间；时间允许扩展至10个",
        "问题三只画流程图": "问题三只画流程图或试图一次做全四类事件",
        "真实模拟四类事件": "至少真实模拟2类代表事件，其余事件给出统一响应流程",
    }
    for old, new in replacements.items():
        replace_everywhere(doc, old, new)

    # Rebalance role 1: solver and integration, not the entire data pipeline.
    replace_exact_paragraph(
        doc,
        "完成 data_loader.py、preprocess.py 和 evaluator.py。",
        "定义 DataBundle、车辆、任务与解结构接口；和2号共同完成 evaluator.py，不重复维护数据清洗逻辑。",
    )
    replace_exact_paragraph(
        doc,
        "三天个人交付：可运行主程序、统一评估器、问题一至三求解脚本、requirements.txt、运行说明。",
        "三天个人交付：可运行主程序、问题一ALNS、问题二限行版本、问题三局部重调度、requirements.txt与README。",
    )

    # Role 2 owns the specification, reference calculation, and acceptance tests.
    replace_exact_paragraph(
        doc,
        "与1号逐行确认 evaluator.py 中每个成本项的业务含义。",
        "主责 evaluator.py 的计算口径、参考计算和验收测试；1号负责高性能实现，二者必须逐项一致。",
    )
    replace_exact_paragraph(
        doc,
        "三天个人交付：模型说明、手算测试表、小规模精确解、敏感性分析、模型检验章节。",
        "三天个人交付：模型说明、手算测试表、单元测试、小规模精确解、问题二/三指标和模型章节。",
    )

    # Role 3 owns the one and only processed-data pipeline.
    replace_exact_paragraph(
        doc,
        "完成 data_audit.py/Notebook：行数、列名、缺失值、异常值、需求聚合、绿色区识别。",
        "主责 data_loader.py、preprocess.py 和 data_audit.py/Notebook：缺失值、异常值、需求聚合、绿色区与拆分任务。",
    )
    replace_exact_paragraph(
        doc,
        "设计问题三的轻度、中度、重度动态场景，并保存事件输入文件。",
        "至少设计新增订单和时间窗调整2个可复现实验；取消与地址变更保留统一策略接口。",
    )
    replace_exact_paragraph(
        doc,
        "三天个人交付：数据审计、实验记录、全部论文图表、动态场景文件、论文第一版。",
        "三天个人交付：唯一清洗数据、数据审计、批量实验脚本、核心图表、2个动态场景和论文整合稿。",
    )

    # RACI: each artifact has one clear owner and a separate acceptor.
    raci = find_table(doc, ("工作项", "1号", "2号", "3号"))
    raci_rows = [
        ["工作项", "1号", "2号", "3号"],
        ["数据清洗", "确认接口", "审核口径", "主责"],
        ["模型假设", "实现确认", "主责", "记录"],
        ["路线评估器", "性能实现", "规则+验收", "输出格式"],
        ["ALNS", "主责", "检查目标", "批量实验"],
        ["Gurobi小实例", "配合接口", "主责", "汇总结果"],
        ["问题二", "限行实现", "政策验收", "对比实验"],
        ["问题三", "重调度实现", "指标验收", "场景实验"],
        ["论文", "算法章节", "模型章节", "主编整合"],
    ]
    write_table(raci, raci_rows)

    # First-day schedule now reaches a feasible Q1 solution, not merely a route evaluator.
    first_day = find_table(doc, ("时间", "负责人", "任务", "交付物"))
    first_day_rows = [
        ["时间", "负责人", "任务", "交付物"],
        ["0-0.5小时", "全员", "读题、填角色、冻结单位与数据接口", "项目结构+model_spec草案"],
        ["0.5-3小时", "1号", "定义解结构、评估器接口和人工样例", "代码骨架+sample fixture"],
        ["0.5-3小时", "2号", "整理公式、约束、歧义和手算路线", "model_spec+manual_cases"],
        ["0.5-3小时", "3号", "数据审计、缺失处理、聚合与拆分", "processed数据+audit"],
        ["第3小时", "全员", "确认绿色区、拆分、17时后速度等口径", "model_spec v0.1签字"],
        ["3-6小时", "1号+2号", "实现并验收evaluator.py与单元测试", "手算与程序逐项一致"],
        ["3-6小时", "3号", "建立结果格式、实验登记和论文骨架", "schema+registry+outline"],
        ["6-10小时", "全员并行", "1号构造可行解；2号查可行性；3号出基础图", "问题一首个可行方案"],
        ["第10小时", "全员", "演示评估器和问题一可行方案", "第一质量门通过"],
    ]
    write_table(first_day, first_day_rows)

    replace_exact_paragraph(doc, "4. 七天完整推进计划", "4. 三天冲刺推进计划")
    heading_4 = next(p for p in doc.paragraphs if p.text.strip() == "4. 三天冲刺推进计划")
    heading_4.paragraph_format.page_break_before = True
    plan = find_table(doc, ("日期", "主题", "主要动作", "当天验收"))
    plan_rows = [
        ["日期", "并行分工", "主要动作", "当天验收"],
        ["第1天", "1号解结构；2号模型验收；3号数据管道", "完成数据口径、统一评估器、初始解和局部搜索", "评估器通过；问题一有可行方案"],
        ["第2天", "1号ALNS+限行；2号精确验证；3号批量实验", "中午冻结问题一；晚上完成问题二合规方案", "Q1/Q2结果、对比表、可行性报告"],
        ["第3天", "1号动态重调度；2号总验收；3号场景+论文", "上午完成2类动态事件；14:00冻结代码参数；晚间全文核数", "三问结果齐全、论文与附件可提交"],
    ]
    write_table(plan, plan_rows)
    insert_blocks_after(
        doc,
        plan._tbl,
        [
            ("4.1 三天冲刺取舍原则", "Heading 2"),
            ("必做：统一评估器、三问可行方案、问题一/二公平对比、2类动态事件、完整路线与成本表。", "List Bullet"),
            ("尽量做：5个随机种子、200-500次蒙特卡洛、8与12客户精确验证、2-3个关键参数敏感性。", "List Bullet"),
            ("时间允许再做：10个随机种子、1000次蒙特卡洛、完整消融、15客户以上精确求解和四类动态事件。", "List Bullet"),
            ("主动放弃：30客户精确最优、大规模算法横向堆叠、非核心美化和无法回溯到结果文件的额外图表。", "List Bullet"),
        ],
    )
    replace_exact_paragraph(doc, "4.1 每天固定节奏", "4.2 每天固定节奏")

    # The source forced section 7.3 onto a new page, leaving the preceding page mostly blank.
    heading_73 = next(p for p in doc.paragraphs if p.text.strip() == "7.3 代码合并规则")
    heading_73.paragraph_format.page_break_before = False

    replace_exact_paragraph(doc, "9. 实验设计：做到“很好”必须补上的证据", "9. 三天实验设计：最小充分证据")
    replace_exact_paragraph(doc, "9.2 建议敏感性参数", "9.2 精简敏感性参数")
    heading_92 = next(p for p in doc.paragraphs if p.text.strip() == "9.2 精简敏感性参数")
    insert_blocks_after(
        doc,
        heading_92._p,
        [("三天内只选2-3个最能支撑结论的参数；其余保留为时间允许项。", "Normal")],
    )

    comparison = find_table(doc, ("方法", "作用", "报告指标"))
    comparison_rows = [
        ["方法", "作用", "报告指标"],
        ["贪心插入", "基础可行性", "成本、车辆数、运行时间"],
        ["贪心+局部搜索", "说明局部搜索效果", "相对改进率"],
        ["ALNS（5种子）", "最终主算法", "最好、均值、标准差、收敛"],
        ["Gurobi小规模", "8/12客户最优参照", "最优性差距和时间"],
        ["消融ALNS（可选）", "时间允许再验证算子贡献", "去掉核心算子的成本变化"],
    ]
    write_table(comparison, comparison_rows)

    # Make the dynamic experiment scope explicit while preserving all four response strategies.
    replace_everywhere(
        doc,
        "设计问题三的轻度、中度、重度动态场景",
        "设计问题三的2个代表性动态场景",
    )

    # Final manual wording for three-day evidence and cutoffs.
    replace_everywhere(doc, "第7天", "第3天")
    replace_everywhere(doc, "四类事件 + 局部滚动优化", "2类代表事件 + 局部滚动优化")
    replace_everywhere(doc, "问题三实验结果", "问题三2类事件实验结果")
    replace_everywhere(doc, "500-1000", "200-500")
    replace_everywhere(doc, "至少10个随机种子", "至少5个随机种子")

    dst.parent.mkdir(parents=True, exist_ok=True)
    doc.save(dst)
    print(f"Saved: {dst}")


if __name__ == "__main__":
    if len(sys.argv) != 3:
        raise SystemExit("Usage: edit_team_manual_3day.py INPUT.docx OUTPUT.docx")
    main(Path(sys.argv[1]), Path(sys.argv[2]))
